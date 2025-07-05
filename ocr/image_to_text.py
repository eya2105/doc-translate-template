import cv2
import pytesseract
import numpy as np
from typing import List, Dict, Any, Tuple
from spellchecker import SpellChecker
import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from pdf2image import convert_from_path
except ImportError:
    convert_from_path = None

class ImageToText:
    def __init__(self):
        self.bullet_symbols = {
            'circle': '•',
            'square': '■',
            'check': '✓',
            'cross': '✗'
        }
        # Initialize spell checker
        self.spell = SpellChecker()
        
        # Common meaningless patterns for OCR errors
        self.meaningless_patterns = [
            r'^([a-zA-Z])\1{2,}$', # Any letter repeated 3+ times
            r'^[^a-zA-Z0-9\s]+$', # Only special characters
            r'^[0-9]+[a-zA-Z]{1,2}$', # Numbers followed by 1-2 letters
            r'^[a-zA-Z]{1,2}[0-9]+$', # 1-2 letters followed by numbers
        ]
        

    def _preprocess_image(self, img: np.ndarray) -> np.ndarray:
        """Basic image preprocessing for OCR"""
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
        return cv2.GaussianBlur(thresh, (3, 3), 0)

    def _is_meaningless_word(self, word: str) -> bool:
        """Check if a word is likely a meaningless OCR error"""
        if not word or len(word) < 2:
            return False
            
        # Check against meaningless patterns
        for pattern in self.meaningless_patterns:
            if re.match(pattern, word):
                return True
        
        # Check for words with unusual character repetition
        if len(word) >= 3:
            char_counts = {}
            for char in word.lower():
                char_counts[char] = char_counts.get(char, 0) + 1
            
            # If any character appears more than 50% of the time and word is short
            max_char_ratio = max(char_counts.values()) / len(word)
            if max_char_ratio > 0.5 and len(word) <= 5:
                return True
        
        return False

    def _filter_nonsense_words(self, text: str) -> str:
        """
        Enhanced filtering focusing on short meaningless words and OCR errors
        """
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Skip empty lines
            if not line.strip():
                cleaned_lines.append(line)
                continue
                
            # Split into words and non-words (punctuation, spaces)
            tokens = re.findall(r'(\w+|\W+)', line)
            cleaned_tokens = []
            
            for token in tokens:
                # Keep non-word tokens as-is (spaces, punctuation)
                if not token.strip() or not re.match(r'\w+', token):
                    cleaned_tokens.append(token)
                    continue
                
                # Check if word should be kept
                lower_token = token.lower()
                
                
                # Check if it's a meaningless word
                if self._is_meaningless_word(token):
                    cleaned_tokens.append('')  # Remove meaningless words
                    continue
                
                # For other words, apply original logic
                is_valid = (
                    len(token) >= 3 or  # Keep words 3+ characters
                    token.istitle() or  # proper nouns
                    token.isupper() or  # acronyms
                    token in self.spell or  # in dictionary
                    any(char.isdigit() for char in token) or  # alphanumeric
                    (len(token) > 2 and token.lower() in self.spell)  # check lowercase
                )
                
                if is_valid:
                    cleaned_tokens.append(token)
                else:
                    # Replace nonsense words with empty string
                    cleaned_tokens.append('')
                    
            # Reconstruct the line while cleaning up extra spaces
            cleaned_line = ''.join(cleaned_tokens)
            # Clean up multiple spaces
            cleaned_line = re.sub(r'\s+', ' ', cleaned_line)
            cleaned_lines.append(cleaned_line)
            
        return '\n'.join(cleaned_lines)

    def _get_text_boxes(self, img: np.ndarray) -> List[Tuple[int, int, int, int]]:
        """Helper to get text boxes from pytesseract"""
        try:
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            data = pytesseract.image_to_data(gray, output_type=pytesseract.Output.DICT, config='--oem 3 --psm 6')
            return [(data['left'][i], data['top'][i], 
                     data['left'][i] + data['width'][i], 
                     data['top'][i] + data['height'][i]) 
                    for i, word in enumerate(data['text']) if word.strip()]
        except Exception:
            return []

    def _overlaps_text(self, box: Tuple[int, int, int, int], 
                      text_boxes: List[Tuple[int, int, int, int]], 
                      threshold: float = 0.3) -> bool:
        """Check if box overlaps with text regions"""
        x1, y1, x2, y2 = box
        for tx1, ty1, tx2, ty2 in text_boxes:
            ix1, iy1 = max(x1, tx1), max(y1, ty1)
            ix2, iy2 = min(x2, tx2), min(y2, ty2)
            iw, ih = max(0, ix2 - ix1), max(0, iy2 - iy1)
            if iw * ih > threshold * (x2 - x1) * (y2 - y1):
                return True
        return False

    def _filter_overlapping_boxes(self, boxes: List[Dict[str, Any]], max_items: int = 2) -> List[Dict[str, Any]]:
        """Filter overlapping boxes keeping the largest ones"""
        def overlap(a, b):
            return not (a[2] < b[0] or a[0] > b[2] or a[3] < b[1] or a[1] > b[3])

        filtered = []
        for i, item in enumerate(boxes):
            keep = True
            for j, other in enumerate(boxes):
                if i != j and overlap(item['box'], other['box']):
                    a_area = (item['box'][2]-item['box'][0])*(item['box'][3]-item['box'][1])
                    b_area = (other['box'][2]-other['box'][0])*(other['box'][3]-other['box'][1])
                    if a_area < b_area:
                        keep = False
                        break
            if keep and (max_items is None or len(filtered) < max_items):
                filtered.append(item)
        return filtered

    def detect_bullet_points(self, img: np.ndarray) -> List[Dict[str, Any]]:
        """Detect circular, square, check, and cross bullet points"""
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        blurred = cv2.medianBlur(gray, 5)
        bullet_boxes = []

        # Detect circles
        circles = cv2.HoughCircles(
            blurred, cv2.HOUGH_GRADIENT, dp=1.2, minDist=20,
            param1=50, param2=30, minRadius=7, maxRadius=20
        )
        if circles is not None:
            circles = np.round(circles[0, :]).astype("int")
            for (x, y, r) in circles:
                area = np.pi * r * r
                if area > 120:  # Filter out small circles
                    bullet_boxes.append({'box': [x - r, y - r, x + r, y + r], 'type': 'circle'})

        # Detect squares using contours
        _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)
        contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        for cnt in contours:
            approx = cv2.approxPolyDP(cnt, 0.04 * cv2.arcLength(cnt, True), True)
            area = cv2.contourArea(cnt)
            if len(approx) == 4 and area > 60 and cv2.isContourConvex(approx):
                x, y, w, h = cv2.boundingRect(approx)
                aspect = w / h
                if 0.8 < aspect < 1.2 and w > 10 and h > 10:
                    bullet_boxes.append({'box': [x, y, x + w, y + h], 'type': 'square'})

        # Detect check marks and crosses using template matching
        templates = [
            ('check', np.array([[0,0,1,0,0], [0,1,0,0,0], [1,0,0,0,0], [0,1,0,0,0], [0,0,1,0,0]], dtype=np.uint8) * 255),
            ('cross', np.array([[1,0,0,0,1], [0,1,0,1,0], [0,0,1,0,0], [0,1,0,1,0], [1,0,0,0,1]], dtype=np.uint8) * 255)
        ]
        for typ, template in templates:
            for scale in [15, 20, 25]:
                tpl = cv2.resize(template, (scale, scale), interpolation=cv2.INTER_NEAREST)
                res = cv2.matchTemplate(thresh, tpl, cv2.TM_CCOEFF_NORMED)
                loc = np.where(res > 0.7)
                for pt in zip(*loc[::-1]):
                    x, y = pt
                    if scale > 10:
                        bullet_boxes.append({'box': [x, y, x + scale, y + scale], 'type': typ})

        # Remove boxes that overlap with text boxes
        text_boxes = self._get_text_boxes(img)
        filtered = []
        for item in bullet_boxes:
            if not self._overlaps_text(tuple(item['box']), text_boxes, threshold=0.5):
                filtered.append(item)
        return self._filter_overlapping_boxes(filtered, max_items=None)

    def _insert_bullet_points(self, text: str, img: np.ndarray, bullet_items: List[Dict[str, Any]]) -> str:
        """Insert bullet point symbols into the text"""
        lines = text.splitlines()
        h, _ = img.shape[:2]
        
        # Get line positions from OCR data or estimate
        try:
            processed = self._preprocess_image(img)
            data = pytesseract.image_to_data(processed, output_type=pytesseract.Output.DICT, config='--oem 3 --psm 6')
            line_positions = []
            for i, word in enumerate(data['text']):
                if word.strip() and (not line_positions or abs(data['top'][i] - line_positions[-1]) > 10):
                    line_positions.append(data['top'][i])
        except Exception:
            line_positions = [int(h * i / max(1, len(lines))) for i in range(len(lines))]

        # Map bullets to lines
        bullet_lines = {}
        for item in bullet_items:
            x1, y1, x2, y2 = item['box']
            bullet_y = (y1 + y2) // 2
            closest_line = min(enumerate(line_positions), key=lambda x: abs(x[1] - bullet_y), default=(0, 0))[0]
            bullet_lines[closest_line] = item['type']

        # Insert bullet symbols
        for idx, typ in bullet_lines.items():
            if 0 <= idx < len(lines):
                symbol = self.bullet_symbols.get(typ, '•')
                if not lines[idx].strip().startswith(symbol):
                    lines[idx] = f"{symbol} " + lines[idx].lstrip()

        return "\n".join(lines)

    def _extract_with_layout(self, img: np.ndarray) -> str:
        """Enhanced version using hOCR for better layout preservation"""
        try:
            # Get hOCR output for spatial data
            hocr = pytesseract.image_to_pdf_or_hocr(
                self._preprocess_image(img),
                extension='hocr',
                config='--psm 6 --oem 3 -c preserve_interword_spaces=1'
            )
            
            # Parse with BeautifulSoup
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(hocr, 'html.parser')
            
            # Reconstruct layout
            output = []
            for para in soup.find_all(class_='ocr_par'):
                para_text = []
                for line in para.find_all(class_='ocr_line'):
                    line_text = ' '.join(word.get_text() for word in line.find_all(class_='ocrx_word'))
                    if line_text.strip():
                        # Add indentation detection
                        try:
                            x_coord = int(line['title'].split()[1])
                            indent = '    ' * (x_coord // 100)  # 100px per indent level
                            para_text.append(f"{indent}{line_text}")
                        except (IndexError, ValueError):
                            para_text.append(line_text)
                if para_text:
                    output.append('\n'.join(para_text))
            
            return '\n\n'.join(output)
        except Exception:
            # Fallback to basic extraction
            return pytesseract.image_to_string(self._preprocess_image(img), config='--oem 3 --psm 6')

    def save_to_docx(self, text: str, output_path: str, title: str = "OCR Extracted Text"):
        """Save the extracted text to a DOCX file with proper formatting"""
        doc = Document()
        
        # Add title
        title_paragraph = doc.add_heading(title, 0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add extracted text
        paragraphs = text.split('\n\n')
        for paragraph_text in paragraphs:
            if paragraph_text.strip():
                lines = paragraph_text.split('\n')
                for line in lines:
                    if line.strip():
                        p = doc.add_paragraph()
                        
                        # Handle bullet points
                        if line.strip().startswith(('•', '■', '✓', '✗')):
                            p.style = 'List Bullet'
                            p.add_run(line.strip()[1:].strip())  # Remove bullet symbol
                        # Handle indented text
                        elif line.startswith('    '):
                            p.paragraph_format.left_indent = Inches(0.5)
                            p.add_run(line.strip())
                        else:
                            p.add_run(line.strip())
        
        # Save the document
        doc.save(output_path)
        print(f"✅ Document saved to: {output_path}")

    def extract_text(self, image_path: str) -> Dict[str, Any]:
        """Main extraction method. Supports image files and PDFs (first page)."""
        # Detect if PDF
        if image_path.lower().endswith('.pdf'):
            if convert_from_path is None:
                raise ImportError("pdf2image is required for PDF support. Install with 'pip install pdf2image'.")
            pages = convert_from_path(image_path, dpi=150, first_page=1, last_page=1)
            if not pages:
                raise FileNotFoundError(f"No pages found in PDF '{image_path}'.")
            img = np.array(pages[0])
            if img.shape[2] == 4:  # RGBA to RGB
                img = cv2.cvtColor(img, cv2.COLOR_RGBA2RGB)
        else:
            img = cv2.imread(image_path)
            if img is None:
                raise FileNotFoundError(f"Image file '{image_path}' not found.")

        # Extract text with layout preservation
        text = self._extract_with_layout(img)

        # Filter out nonsense words
        text = self._filter_nonsense_words(text)

        # Detect bullet points
        bullet_items = self.detect_bullet_points(img)
        text_with_bullets = self._insert_bullet_points(text, img, bullet_items)

        return {
            "text": text_with_bullets.strip(),
            "bullets": [item['box'] for item in bullet_items]
        }

    def process_image(self, image_path: str, output_docx: str = None) -> Dict[str, Any]:
        """Process image and optionally save to DOCX"""
        result = self.extract_text(image_path)
        
        # Save to DOCX if output path provided
        if output_docx:
            self.save_to_docx(result["text"], output_docx, f"OCR Results - {image_path}")
        
        return result

if __name__ == "__main__":
    ocr = ImageToText()
    
    # Process the image/PDF
    result = ocr.process_image('Vollmacht  (2).pdf', 'processed_document.docx')
    
    print("🔍 Extracted Text:\n", result["text"])
    print("📄 Detected bullet points:", result["bullets"])
    print("💾 Text also saved to 'processed_document.docx'")
    
    # Still save to txt for backup
    with open("output.txt", "w", encoding="utf-8") as f:
        f.write(result["text"])